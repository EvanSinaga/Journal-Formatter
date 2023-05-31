from docx import Document
from docx.shared import Cm, Pt
from readDocx import *
from docx.oxml.ns import qn
from pageNumber import add_page_number

import pymysql

# database connection
connection = pymysql.connect(
    host="localhost", user="root", passwd="", database="sasjep")
cursor = connection.cursor()

# # query to retrieve the newly inserted file
# retrievefile = "SELECT file FROM artikel WHERE id= (SELECT MAX(id) FROM artikel);"

# # executing the query
# cursor.execute(retrievefile)
# fileresult = cursor.fetchone()

# # extracting the file name
# initialfile = str(fileresult)
# editfile = initialfile.replace("('", "")
# filename = editfile.replace("',)", "")

"""
retrieve = "SELECT year, volume, number, file FROM artikel WHERE id= (SELECT MAX(id) FROM artikel);"
cursor.execute(retrieve)
retrieveOut = cursor.fetchone() #('year','volume','number','file')

# ('year','volume','number','file')
# print(retrieveOut[0]) -> year

year = retrieveOut[0]
volume = retrieveOut[1]
number = retrieveOut[2]
file = retrieveOut[3]
"""
retrieve = "SELECT year, volume, number, namafile, id FROM artikel WHERE id= (SELECT MAX(id) FROM artikel);"
cursor.execute(retrieve)
retrieveOut = cursor.fetchone()

year = retrieveOut[0]
volume = retrieveOut[1]
number = retrieveOut[2]
filename = retrieveOut[3]
id_jurnal = str(retrieveOut[4])

# processing the file
input_doc = Document("C:/xampp/htdocs/pkl/assets/temp/" + filename)

# Input Volume, nomor, tahun
htext1 = 'IZUMI, Volume '+volume+' No. '+number + ', '+year+', '
htext2 = '\ne-ISSN: 2502-3535, p-ISSN: 2338-249X'
htext3 = '\nAvailable online at: http://ejournal.undip.ac.id/index.php/izumi'

# Document margin layout
sections = input_doc.sections
for section in sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Input text for header and footer


def fhd(position, text, nth):
    x = position.add_run(text)
    x.font.name = 'Times New Roman'
    x.font.size = Pt(11)
    position.alignment = nth


# First Header
first_header = input_doc.sections[0].first_page_header
input_doc.sections[0].different_first_page_header_footer = True
# delete first paragraph
delete_paragraph(first_header.paragraphs[0])

# First Header Table
htable = first_header.add_table(1, 2, width=Cm(15))
htable.alignment = 1
htab_cells = htable.rows[0].cells

fht0 = htab_cells[0].paragraphs[0]
htab_cells[0].width = Cm(12)


fhd(fht0, htext1, 2)
add_page_number(fht0)
fhd(fht0, htext2, 2)
fhd(fht0, htext3, 2)

# Gambar
ht1 = htab_cells[1].paragraphs[0]
img1 = ht1.add_run()
img1.add_picture(
    "C:/xampp/htdocs/pkl/assets/assets/Izumi.png", width=Cm(1.5))
ht1.alignment = 3  # Justify

htab_cells[1].width = Cm(3)
# First Footer
first_footer = input_doc.sections[0].first_page_footer
# delete first paragraph
delete_paragraph(first_footer.paragraphs[0])
ftext = 'Copyright@'+year+', IZUMI, e-ISSN: 2502-3535, p-ISSN: 2338-249x'
fhd(first_footer.add_paragraph(), ftext, 1)  # First footer

# Header
header = input_doc.sections[0].header
# delete first paragraph
delete_paragraph(header.paragraphs[0])
htable = header.add_table(1, 1, width=Cm(15))
htable.alignment = 1
htab_cells = htable.rows[0].cells

# Header Content
ht0 = htab_cells[0].paragraphs[0]
fhd(ht0, htext1, 2)
add_page_number(ht0)
fhd(ht0, htext2, 2)
fhd(ht0, htext3, 2)

# Footer
footer = input_doc.sections[0].footer
# delete first paragraph
delete_paragraph(footer.paragraphs[0])
ftext = 'Copyright@'+year+', IZUMI, e-ISSN: 2502-3535, p-ISSN: 2338-249x'
fhd(footer.add_paragraph(), ftext, 1)  # Footer

# Body
section = input_doc.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')

# Mirror
input_doc.save('C:/xampp/htdocs/pkl/assets/file/' + filename)

# .replace('.docx', '') + '-converted.docx'

# INSERT data from python to mysql
# database connection
title = get_title(input_doc)
author = get_author(input_doc)
keywords = get_keywords(input_doc)
cursor.execute(
    "UPDATE artikel SET judul='"+title+"', author='"+author+"', keywords='"+keywords+"'  WHERE id="+id_jurnal+"")
connection.commit()

# commiting the connection then closing it.
cursor.close()
connection.close()
