from docx import Document
# def get_para_data(out_doc_name, paragraph):
#     out_para = out_doc_name.add_paragraph()
#     for run in paragraph.runs:
#         out_run = out_para.add_run(run.text)
#         # Styling
#         out_run.bold = run.bold
#         out_run.italic = run.italic
#         out_run.underline = run.underline
#         out_run.font.color.rgb = run.font.color.rgb
#         out_run.style.name = run.style.name
#     # Paragraph alignment
#     out_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

# Judul


def get_title(input_doc):
    title = input_doc.paragraphs[0].text
    return title

# Author e.g. : Namikaze Minato, Senju Hashirama, Ichinose Mizuhara


def get_author(input_doc):
    i = 0
    for run in input_doc.paragraphs[1].text:
        i += 1
    # print(i)
    if(i != 0):
        return input_doc.paragraphs[1].text
    else:
        return input_doc.paragraphs[2].text

# Keywords e.g. : KNN, Machine Learning, CNN


def get_keywords(input_doc):
    i = 0
    for para in input_doc.paragraphs:
        keywords = input_doc.paragraphs[i].text
        i += 1
        if("Keywords" in para.text):
            keywords = keywords.replace('Keywords: ', '')
            keywords = keywords.replace('Keywords : ', '')
            return keywords


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# def getText(doc):
#     allText = []

#     for paragraph in doc.paragraphs:
#         allText.append(paragraph.text)
#     return "".join(allText).encode("utf-8")

# Mengambil seluruh data docx

# def getTitle():
#     # nPrg = 0
#     # title = []
#     i = 0
#     while i <= 999:
#         prg = doc.paragraphs[i].text
#         if("Abstract" in prg):
#             # print("Total Paragraf = ", nPrg)
#             break
#         print(prg)
#         # nPrg += 1
#         i += 1


# def getAbstract():
#     i = 0
#     while i <= 999:
#         prg = doc.paragraphs[i].text
#         if("Introduction" in prg):
#             break
#         print(prg)
#         i += 1

# getTitle()
# getAbstract()


# print(getText(doc))
